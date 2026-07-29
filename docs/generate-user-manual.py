from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Derek-Control-Panel-使用說明書.md"
OUTPUT = ROOT / "Derek-Control-Panel-使用說明書.docx"

FONT = "Microsoft JhengHei"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "64748B"
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
USABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color, before, after, line):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.add_run("第 ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_char1)
    paragraph._p.append(instr_text)
    paragraph._p.append(fld_char2)
    paragraph.add_run(" 頁")


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11, INK, 0, 6, 1.25)
    for name, size, color, before, after in (
        ("Title", 28, DARK_BLUE, 0, 8),
        ("Subtitle", 12, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        set_style_font(style, size, color, before, after, 1.25)
        if name.startswith("Heading"):
            style.font.bold = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        set_style_font(style, 11, INK, 0, 4, 1.25)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    if "Callout" not in [s.name for s in doc.styles]:
        callout = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = doc.styles["Callout"]
    set_style_font(callout, 11, INK, 6, 6, 1.25)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("Derek Control Panel  |  使用說明書")
    set_run_font(header_run, 8.5, MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("2026.07  |  ")
    set_run_font(footer_run, 8.5, MUTED)
    add_page_field(footer)
    for run in footer.runs:
        set_run_font(run, 8.5, MUTED)


def add_title_block(doc, title, subtitle_lines):
    doc.add_paragraph().paragraph_format.space_after = Pt(64)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, 28, DARK_BLUE, bold=True)
    for line in subtitle_lines:
        sp = doc.add_paragraph(style="Subtitle")
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sp.add_run(line)
        set_run_font(run, 12, MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(42)
    box = doc.add_table(rows=1, cols=1)
    set_table_geometry(box, [USABLE_WIDTH_DXA])
    cell = box.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.style = "Callout"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("不是要同時做更多；而是把下一步變得清楚，讓中斷後更容易回到工作。")
    set_run_font(run, 12, INK, bold=True)
    doc.add_page_break()


def add_heading(doc, level, text):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, {1: 16, 2: 13, 3: 12}[level], {1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)


def add_text_runs(paragraph, text):
    parts = re.split(r"(「[^」]+」|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("「") or part.startswith("**")
        value = part[2:-2] if part.startswith("**") else part
        run = paragraph.add_run(value)
        set_run_font(run, 11, INK, bold=bold)


def add_plain(doc, text):
    p = doc.add_paragraph(style="Normal")
    add_text_runs(p, text)


def add_list(doc, text, numbered):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    add_text_runs(p, text)


def parse_table(lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped or re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", stripped):
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            rows.append([part.strip() for part in stripped[1:-1].split("|")])
    return rows


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    if cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [2100, 2520, 4740]
    else:
        base = USABLE_WIDTH_DXA // cols
        widths = [base] * (cols - 1) + [USABLE_WIDTH_DXA - base * (cols - 1)]
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_text_runs(p, value)
            for run in p.runs:
                set_run_font(run, 10.2, INK, bold=(r_idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(5)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [USABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.style = "Callout"
    add_text_runs(p, text)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def build_docx():
    doc = Document()
    setup_document(doc)
    text = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    title_done = False
    while i < len(text):
        line = text[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            subtitle = []
            j = i + 1
            while j < len(text) and (not text[j].strip() or text[j].lstrip().startswith(">")):
                if text[j].lstrip().startswith(">"):
                    subtitle.append(text[j].split(">", 1)[1].strip())
                j += 1
            if not title_done:
                add_title_block(doc, title, subtitle)
                title_done = True
            i = j
            continue
        match = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if match:
            add_heading(doc, len(match.group(1)) - 1, match.group(2))
            i += 1
            continue
        if stripped.startswith("> "):
            add_callout(doc, stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(text) and text[i].strip().startswith("|"):
                table_lines.append(text[i])
                i += 1
            rows = parse_table(table_lines)
            if rows:
                add_table(doc, rows)
            continue
        num = re.match(r"^\d+\.\s+(.*)$", stripped)
        if num:
            add_list(doc, num.group(1), numbered=True)
            i += 1
            continue
        if stripped.startswith("- "):
            add_list(doc, stripped[2:].strip(), numbered=False)
            i += 1
            continue
        add_plain(doc, stripped.replace("  ", " "))
        i += 1

    props = doc.core_properties
    props.title = "Derek Control Panel 使用說明書"
    props.subject = "Derek 與 Suki 的雙人個人工作系統"
    props.author = "Derek Control Panel"
    props.keywords = "Derek Control Panel, 使用說明書, 任務, 交接, Focus"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
